"""Agent tools for the FleetOps agent.

Three tools, each with a JSON schema (TOOL_SPECS) shared by the agent loop,
the eval harness, and the dashboard:

  * run_sql        — READ-ONLY SQL over fleet.db (SELECT/CTE only, row-capped)
  * lookup_part    — part details + stock status for a fault code
  * generate_report — maintenance summary report (docx, markdown fallback)

Security model for run_sql (defense in depth, all layers must pass):
  1. The database is opened with SQLite's read-only URI mode (mode=ro).
  2. An sqlite3 authorizer callback denies every operation except reading
     table/column data and running built-in functions.
  3. The statement text is validated: single statement, must start with
     SELECT or WITH, and write/DDL keywords are rejected outright.

Tool-calling architecture pattern adapted from LGDiMaggio/predictive-maintenance-mcp
(MIT) and patchy631/ai-engineering-hub rag-sql-router (MIT) — reimplemented here,
no code copied verbatim.
"""

from __future__ import annotations

import re
import sqlite3
from datetime import date, timedelta
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DB = PROJECT_ROOT / "fleet.db"
REPORTS_DIR = PROJECT_ROOT / "reports"
MAX_ROWS = 50

_FORBIDDEN = re.compile(
    r"\b(insert|update|delete|drop|alter|create|replace|attach|detach|"
    r"pragma|vacuum|reindex|trigger|begin|commit|rollback)\b",
    re.IGNORECASE,
)


def _strip_sql_comments(sql: str) -> str:
    sql = re.sub(r"--[^\n]*", " ", sql)
    sql = re.sub(r"/\*.*?\*/", " ", sql, flags=re.DOTALL)
    return sql.strip()


def _authorizer(action: int, arg1, arg2, db_name, trigger) -> int:
    """sqlite3 authorizer: allow reads and built-in functions, deny all else."""
    allowed = {sqlite3.SQLITE_READ, sqlite3.SQLITE_SELECT, sqlite3.SQLITE_FUNCTION}
    return sqlite3.SQLITE_OK if action in allowed else sqlite3.SQLITE_DENY


def run_sql(query: str, db_path: str | Path = DEFAULT_DB) -> dict:
    """Execute a READ-ONLY SQL SELECT against fleet.db.

    Args:
        query: A single SELECT (or WITH ... SELECT) statement. Anything else —
            writes, DDL, PRAGMA, multiple statements — is refused.
        db_path: Path to the SQLite database (default: project fleet.db).

    Returns:
        dict with either:
          {"columns": [...], "rows": [[...], ...], "row_count": int,
           "truncated": bool}  on success, or
          {"error": "<reason>"}  on refusal/failure. Refusals never touch the DB.
    """
    cleaned = _strip_sql_comments(query)
    if not cleaned:
        return {"error": "Refused: empty query."}
    # single statement only (allow one trailing semicolon)
    body = cleaned.rstrip().rstrip(";")
    if ";" in body:
        return {"error": "Refused: multiple SQL statements are not allowed."}
    if not re.match(r"^(select|with)\b", body, re.IGNORECASE):
        return {"error": "Refused: only SELECT queries are allowed (read-only tool)."}
    if _FORBIDDEN.search(body):
        return {"error": "Refused: query contains a write/DDL keyword; this tool is read-only."}

    try:
        con = sqlite3.connect(f"file:{Path(db_path)}?mode=ro", uri=True)
        con.set_authorizer(_authorizer)
        cur = con.execute(body)
        columns = [d[0] for d in cur.description] if cur.description else []
        rows = cur.fetchmany(MAX_ROWS + 1)
        con.close()
    except sqlite3.Error as exc:
        return {"error": f"SQL error: {exc}"}

    truncated = len(rows) > MAX_ROWS
    rows = rows[:MAX_ROWS]
    return {
        "columns": columns,
        "rows": [list(r) for r in rows],
        "row_count": len(rows),
        "truncated": truncated,
    }


def lookup_part(fault_code: str, db_path: str | Path = DEFAULT_DB) -> dict:
    """Look up the recommended part and current stock for a fault code.

    Args:
        fault_code: J1939-style code, e.g. 'SPN-3226-FMI-4'. Case-insensitive.

    Returns:
        dict with fault details, the related part, stock level, and a
        stock_status of 'in_stock' | 'low_stock' | 'out_of_stock',
        or {"error": ...} if the code is unknown or has no related part.
    """
    con = sqlite3.connect(f"file:{Path(db_path)}?mode=ro", uri=True)
    row = con.execute(
        """SELECT f.code, f.description, f.severity, f.system, f.recommended_action,
                  p.part_id, p.part_name, p.part_number, p.unit_cost,
                  p.stock_qty, p.min_stock, p.supplier, p.lead_time_days
           FROM fault_codes f LEFT JOIN parts p ON p.part_id = f.related_part_id
           WHERE UPPER(f.code) = UPPER(?)""",
        (fault_code.strip(),),
    ).fetchone()
    con.close()
    if row is None:
        return {"error": f"Unknown fault code: {fault_code!r}. "
                         "Use run_sql on the fault_codes table to list valid codes."}
    (code, desc, sev, system, action, part_id, name, number, cost,
     stock, min_stock, supplier, lead) = row
    if part_id is None:
        return {"fault_code": code, "description": desc, "severity": sev,
                "system": system, "recommended_action": action, "part": None}
    status = "out_of_stock" if stock == 0 else ("low_stock" if stock < min_stock else "in_stock")
    return {
        "fault_code": code, "description": desc, "severity": sev, "system": system,
        "recommended_action": action,
        "part": {"part_id": part_id, "part_name": name, "part_number": number,
                 "unit_cost": cost, "stock_qty": stock, "min_stock": min_stock,
                 "stock_status": status, "supplier": supplier,
                 "lead_time_days": lead},
    }


def generate_report(truck_id: str | None = None, days: int = 30,
                    db_path: str | Path = DEFAULT_DB,
                    reports_dir: str | Path = REPORTS_DIR) -> dict:
    """Generate a maintenance summary report (docx if python-docx is installed,
    markdown otherwise) for the whole fleet or a single truck.

    Args:
        truck_id: Optional truck, e.g. 'T-042'. None = whole fleet.
        days: Lookback window in days (default 30) ending at the data anchor
            date (2026-07-08, the synthetic dataset's "today").

    Returns:
        dict with report file path, format, and the summary stats included,
        or {"error": ...} for an unknown truck.
    """
    anchor = date(2026, 7, 8)
    since = (anchor - timedelta(days=days)).isoformat()
    con = sqlite3.connect(f"file:{Path(db_path)}?mode=ro", uri=True)

    scope = "fleet"
    if truck_id:
        truck = con.execute("SELECT truck_id, make, model, year, mileage, status"
                            " FROM trucks WHERE UPPER(truck_id)=UPPER(?)",
                            (truck_id.strip(),)).fetchone()
        if truck is None:
            con.close()
            return {"error": f"Unknown truck_id: {truck_id!r}"}
        scope = truck[0]

    where = "WHERE event_date >= ?"
    params: list = [since]
    if truck_id:
        where += " AND UPPER(truck_id) = UPPER(?)"
        params.append(truck_id.strip())

    total, cost, hours = con.execute(
        f"SELECT COUNT(*), COALESCE(SUM(cost),0), COALESCE(SUM(labor_hours),0)"
        f" FROM maintenance_log {where}", params).fetchone()
    by_type = con.execute(
        f"SELECT event_type, COUNT(*), ROUND(SUM(cost),2) FROM maintenance_log {where}"
        f" GROUP BY event_type ORDER BY 2 DESC", params).fetchall()
    top_faults = con.execute(
        f"SELECT fault_code, COUNT(*) FROM maintenance_log {where}"
        f" AND fault_code IS NOT NULL GROUP BY fault_code ORDER BY 2 DESC LIMIT 5",
        params).fetchall()
    events = con.execute(
        f"SELECT event_date, truck_id, event_type, COALESCE(fault_code,''),"
        f" description, cost FROM maintenance_log {where}"
        f" ORDER BY event_date DESC LIMIT 25", params).fetchall()
    con.close()

    stats = {"scope": scope, "window_days": days, "since": since,
             "events": total, "total_cost": round(cost, 2),
             "labor_hours": round(hours, 1),
             "by_type": [{"type": t, "count": c, "cost": s} for t, c, s in by_type],
             "top_faults": [{"code": f, "count": c} for f, c in top_faults]}

    reports_dir = Path(reports_dir)
    reports_dir.mkdir(exist_ok=True)
    stem = f"maintenance_{scope.lower().replace(' ', '_')}_{days}d"
    title = f"Maintenance Summary — {scope} — last {days} days (as of {anchor})"

    try:
        from docx import Document  # python-docx

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Events: {total}   Total cost: ${stats['total_cost']:,}   "
                          f"Labor hours: {stats['labor_hours']}")
        doc.add_heading("Events by type", level=2)
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(("Type", "Count", "Cost ($)")):
            t.rows[0].cells[i].text = h
        for row in by_type:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
        if top_faults:
            doc.add_heading("Top fault codes", level=2)
            for code, c in top_faults:
                doc.add_paragraph(f"{code}: {c} occurrences", style="List Bullet")
        doc.add_heading("Recent events (up to 25)", level=2)
        for ev in events:
            doc.add_paragraph(f"{ev[0]}  {ev[1]}  [{ev[2]}] {ev[3]} — {ev[4]} (${ev[5]:,.2f})",
                              style="List Bullet")
        path = reports_dir / f"{stem}.docx"
        doc.save(path)
        fmt = "docx"
    except ImportError:
        lines = [f"# {title}", "",
                 f"- Events: {total}", f"- Total cost: ${stats['total_cost']:,}",
                 f"- Labor hours: {stats['labor_hours']}", "", "## Events by type"]
        lines += [f"- {t_}: {c} (${s:,})" for t_, c, s in by_type]
        if top_faults:
            lines += ["", "## Top fault codes"]
            lines += [f"- {code}: {c}" for code, c in top_faults]
        path = reports_dir / f"{stem}.md"
        path.write_text("\n".join(lines))
        fmt = "markdown"

    return {"report_path": str(path), "format": fmt, "stats": stats}


# ---------------------------------------------------------------------------
# Anthropic tool schemas — shared by agent, evals, and dashboard.
# ---------------------------------------------------------------------------
TOOL_SPECS = [
    {
        "name": "run_sql",
        "description": (
            "Run a READ-ONLY SQL SELECT query against the fleet maintenance SQLite "
            "database (tables: trucks, maintenance_log, fault_codes, parts). "
            "Only a single SELECT/WITH statement is allowed; writes are refused. "
            "Results are capped at 50 rows — use aggregation or LIMIT for large results."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string",
                          "description": "A single SQL SELECT statement."}
            },
            "required": ["query"],
        },
    },
    {
        "name": "lookup_part",
        "description": (
            "Look up a fault code's meaning, severity, recommended action, and the "
            "replacement part with live stock status (in_stock/low_stock/out_of_stock), "
            "unit cost, supplier and lead time. Use for any part/stock/fault-code question."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "fault_code": {"type": "string",
                               "description": "J1939-style code, e.g. 'SPN-3226-FMI-4'."}
            },
            "required": ["fault_code"],
        },
    },
    {
        "name": "generate_report",
        "description": (
            "Generate a downloadable maintenance summary report (docx) for the whole "
            "fleet or one truck over a lookback window. Returns the file path and the "
            "summary stats. Use when the user asks for a report, summary document, or export."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "truck_id": {"type": "string",
                             "description": "Optional truck id, e.g. 'T-042'. Omit for whole fleet."},
                "days": {"type": "integer", "default": 30,
                         "description": "Lookback window in days (default 30)."},
            },
            "required": [],
        },
    },
]

TOOL_FUNCTIONS = {
    "run_sql": run_sql,
    "lookup_part": lookup_part,
    "generate_report": generate_report,
}
